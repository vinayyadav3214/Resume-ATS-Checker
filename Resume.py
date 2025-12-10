import streamlit as st
from docx import Document
from io import BytesIO
import json
import re
import os
from dotenv import load_dotenv
load_dotenv()
from langchain_openai import ChatOpenAI
from langchain_groq import ChatGroq
from langchain_core.messages import HumanMessage, SystemMessage


# ================================
# 1. Extract DOCX Structure (Paragraphs + Styles)
# ================================
def extract_doc_structure(doc_file):
    doc = Document(doc_file)
    structure = []

    for p in doc.paragraphs:
        structure.append({
            "text": p.text,
            "style": p.style.name if p.style else "Normal"
        })

    return structure


# ================================
# 2. Rebuild DOCX With Updated Text
# ================================
def rebuild_docx(structure, replacements):
    doc = Document()

    replace_map = {r["original"]: r["updated"] for r in replacements}

    for block in structure:
        text = block["text"]
        style = block["style"]

        new_text = replace_map.get(text, text)

        para = doc.add_paragraph(style=style)
        para.add_run(new_text)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


# ================================
# 3. Clean Model JSON Output
# ================================
def clean_json(raw_output):
    cleaned = re.sub(r"```.*?```", "", raw_output, flags=re.DOTALL).strip()
    first = cleaned.find("{")
    last = cleaned.rfind("}")
    if first != -1 and last != -1:
        cleaned = cleaned[first:last+1]
    return cleaned


# ================================
# Initialize Groq
# ================================
groq_api_key = os.getenv("GROQ_API_KEY")
llm = ChatGroq(model="llama-3.1-8b-instant", temperature=0)


# ================================
# STREAMLIT UI
# ================================
st.title("ATS Resume Evaluator + Formatter-Preserved Rewriter")
st.write("Upload a DOCX resume and paste the Job Description below.")

uploaded_file = st.file_uploader("Upload Resume (.docx only)", type=["docx", "doc"])
job_description = st.text_area("Paste Job Description", height=250)


if st.button("Evaluate & Rewrite Resume"):

    # --- Validations ---
    if not uploaded_file:
        st.error("Please upload a DOCX resume.")
        st.stop()
    if not job_description.strip():
        st.error("Please paste a Job Description.")
        st.stop()

    # --- Extract original structure ---
    uploaded_file.seek(0)
    structure = extract_doc_structure(uploaded_file)

    # Convert structure into readable text for LLM
    resume_raw_text = "\n".join([b["text"] for b in structure])

    st.subheader("Original Resume Preview")
    st.text_area("", resume_raw_text[:1500], height=200)


    # ================================
    # PHASE 1: ATS EVALUATION
    # ================================
    scoring_rules = """
Calculate ATS score:

1. Keyword Match (40%)
2. Skills Match (20%)
3. Experience Alignment (25%)
4. Resume Structure (15%)

Final score = sum of these components.
Round to nearest integer.
"""

    evaluation_prompt = f"""
    You are an ATS resume evaluator.
    Return ONLY PURE JSON. No markdown. No commentary.

    Use the scoring rules here:
    {scoring_rules}

    JSON format:
    {{
    "ats_score": 0,
    "match_summary": "",
    "missing_keywords": [],
    "missing_skills": [],
    "experience_gaps": [],
    "improvement_suggestions": [],
    "sections_to_update": []
    }}

    --- Resume ---
    {resume_raw_text}

    --- Job Description ---
    {job_description}
    """

    eval_messages = [
        SystemMessage(content="You evaluate resumes for ATS scoring. Return only JSON."),
        HumanMessage(content=evaluation_prompt)
    ]

    with st.spinner("Evaluating Resume..."):
        eval_response = llm.invoke(eval_messages)

    cleaned_eval = clean_json(eval_response.content)

    try:
        eval_json = json.loads(cleaned_eval)
        st.subheader("ATS Evaluation")
        st.json(eval_json)
    except:
        st.error("Could not parse model JSON. Here is raw output:")
        st.text(eval_response.content)
        st.stop()


    # ================================
    # PHASE 2 — SECTION-WISE REWRITE
    # ================================
    replacements = []

    for block in structure:
        original_text = block["text"]

        # Skip empty or header-only lines
        if not original_text.strip():
            continue

        rewrite_prompt = f"""
You are a resume rewriting engine.

Rewrite ONLY the following block of text to align with the job description and ATS evaluation.
Do NOT invent new experience or roles.
Use same length and similar line structure.
Return ONLY JSON:

{{
   "original": "",
   "updated": ""
}}

--- ORIGINAL BLOCK ---
{original_text}

--- JOB DESCRIPTION ---
{job_description}

--- EVALUATION JSON ---
{json.dumps(eval_json)}
"""

        rewrite_messages = [
            SystemMessage(content="Rewrite resume content precisely and concisely."),
            HumanMessage(content=rewrite_prompt)
        ]

        with st.spinner("Rewriting resume content..."):
            rewrite_resp = llm.invoke(rewrite_messages)

        cleaned_rewrite = clean_json(rewrite_resp.content)

        try:
            rewrite_json = json.loads(cleaned_rewrite)
            replacements.append(rewrite_json)
        except:
            # fallback: keep original
            replacements.append({"original": original_text, "updated": original_text})


    # ================================
    # PHASE 3 — Build updated DOCX
    # ================================
    uploaded_file.seek(0)
    output_doc = rebuild_docx(structure, replacements)

    st.success("Resume rewritten successfully!")

    st.download_button(
        "Download Rewritten Resume (.docx)",
        data=output_doc,
        file_name="Rewritten_Resume.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
